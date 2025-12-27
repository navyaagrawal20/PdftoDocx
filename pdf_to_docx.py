"""
PDF to Word Document Converter
This script reads a PDF file and creates an MS Word document replicating its content and layout.
"""

import os
import pdfplumber
from pdf2docx import Converter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_border_to_paragraph(paragraph):
    """Add border to a paragraph"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), '000000')
        pBdr.append(border)
    
    pPr.append(pBdr)

def extract_pdf_content(pdf_path):
    """Extract content from PDF with layout information"""
    content = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                # Split into lines while preserving structure
                lines = text.split('\n')
                content.extend(lines)
    
    return content


# --- New layout-aware extraction helpers ---
def _group_words_by_line(words, tolerance=2):
    """Group words that share the same line based on their top position."""
    lines = []
    if not words:
        return lines

    # Sort by vertical then horizontal position
    words = sorted(words, key=lambda w: (round(w.get("top", 0)), w.get("x0", 0)))
    current_top = round(words[0].get("top", 0))
    current_line = []

    for w in words:
        top = round(w.get("top", 0))
        if abs(top - current_top) <= tolerance:
            current_line.append(w)
        else:
            if current_line:
                lines.append(_build_line_block(current_line))
            current_line = [w]
            current_top = top
    if current_line:
        lines.append(_build_line_block(current_line))
    return lines


def _build_line_block(words):
    """Return metadata for a single text line block."""
    words_sorted = sorted(words, key=lambda x: x.get("x0", 0))
    tops = [w.get("top", 0) for w in words_sorted]
    bottoms = [w.get("bottom", w.get("top", 0)) for w in words_sorted]
    centers = [((w.get("x0", 0) + w.get("x1", 0)) / 2) for w in words_sorted]
    sizes = [w.get("size") for w in words_sorted if w.get("size")]
    fonts = [w.get("fontname", "") for w in words_sorted if w.get("fontname")]
    return {
        "words": words_sorted,
        "top": min(tops) if tops else 0,
        "bottom": max(bottoms) if bottoms else 0,
        "center_x": sum(centers) / len(centers) if centers else 0,
        "font_size": _median(sizes) if sizes else None,
        "font_name": fonts[0] if fonts else None,
        "is_bold": any("bold" in f.lower() or "black" in f.lower() or "heavy" in f.lower() for f in fonts),
    }


def _median(values):
    if not values:
        return None
    values = sorted(values)
    n = len(values)
    mid = n // 2
    if n % 2:
        return values[mid]
    return (values[mid - 1] + values[mid]) / 2


def _add_text_block(doc, line_block, next_top, page_width, default_font="Times New Roman"):
    """Add a single line block with heuristic font, bold, alignment, spacing."""
    text = " ".join(w.get("text", "") for w in line_block.get("words", [])).strip()
    if not text:
        return

    p = doc.add_paragraph()
    run = p.add_run(text)

    # Font size and name
    size = line_block.get("font_size") or 11
    run.font.size = Pt(size)
    if line_block.get("font_name"):
        run.font.name = line_block.get("font_name")
    else:
        run.font.name = default_font

    # Bold detection heuristic
    if line_block.get("is_bold"):
        run.bold = True

    # Alignment heuristic: center if line is visually centered
    center_x = line_block.get("center_x", 0)
    if page_width and abs(center_x - page_width / 2) <= page_width * 0.08:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Spacing heuristic based on gap to next block
    if next_top is not None:
        gap = max(0, next_top - line_block.get("bottom", next_top))
        # PDF units are points; clamp to a reasonable value
        p.paragraph_format.space_after = Pt(min(gap, 12))


def _add_table(doc, table_cells, col_positions=None):
    """Create a docx table from extracted table cells and set column widths if available."""
    if not table_cells:
        return
    rows = len(table_cells)
    cols = max(len(r) for r in table_cells)
    docx_table = doc.add_table(rows=rows, cols=cols)
    docx_table.style = "Table Grid"

    # Fill cells
    for r_idx, row in enumerate(table_cells):
        for c_idx, cell in enumerate(row):
            docx_table.cell(r_idx, c_idx).text = cell if cell is not None else ""

    # Apply column widths based on detected column boundaries (points -> inches)
    if col_positions and len(col_positions) >= cols + 1:
        for c_idx in range(cols):
            width_pts = col_positions[c_idx + 1] - col_positions[c_idx]
            docx_table.columns[c_idx].width = Inches(max(width_pts, 1) / 72)


def _add_page_number_footer(doc):
    """Add dynamic page number field to all section footers."""
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Clear existing text to avoid duplicates
        p.text = "Page "
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'PAGE')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = ''
        r.append(t)
        fld.append(r)
        p._p.append(fld)


def create_word_document_with_layout(pdf_path, output_path):
    """Create Word document preserving text, tables, and approximate layout."""
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    doc = Document()

    # Establish a sane default style to stabilize fonts when PDF metadata is missing
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    table_settings = {
        "vertical_strategy": "lines",
        "horizontal_strategy": "lines",
        "intersection_tolerance": 5,
    }

    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)
        for page_index, page in enumerate(pdf.pages):
            page_width = page.width or 0
            words = page.extract_words(
                extra_attrs=["top", "x0", "bottom", "x1", "fontname", "size"],
                use_text_flow=True,
            )

            # Determine dominant font/size on this page to use as fallback
            sizes = [w.get("size") for w in words if w.get("size")]
            fonts = [w.get("fontname") for w in words if w.get("fontname")]
            dominant_size = _median(sizes) if sizes else 11
            dominant_font = fonts[0] if fonts else "Times New Roman"

            lines = _group_words_by_line(words)

            tables = []
            for tbl in page.find_tables(table_settings=table_settings):
                tables.append({
                    "bbox": tbl.bbox,
                    "cells": tbl.extract(),
                    "cols": getattr(tbl, "cols", None),
                })

            components = []
            for l in lines:
                components.append({
                    "type": "text",
                    "top": l.get("top", 0),
                    "bottom": l.get("bottom", l.get("top", 0)),
                    "data": l,
                })
            for t in tables:
                components.append({
                    "type": "table",
                    "top": t["bbox"][1],
                    "bottom": t["bbox"][3],
                    "data": t["cells"],
                })

            components = sorted(components, key=lambda c: c.get("top", 0))

            for idx, comp in enumerate(components):
                next_top = components[idx + 1]["top"] if idx + 1 < len(components) else None
                if comp["type"] == "text":
                    _add_text_block(doc, comp["data"], next_top, page_width, default_font=dominant_font or "Times New Roman")
                elif comp["type"] == "table":
                    _add_table(doc, comp["data"], col_positions=comp.get("cols"))
                    if next_top is not None:
                        spacer = doc.add_paragraph()
                        spacer.paragraph_format.space_after = Pt(min(max(0, next_top - comp.get("bottom", next_top)), 12))

            # Page break between PDF pages, but not after the last one
            if page_index < total_pages - 1:
                doc.add_page_break()

    # Add page numbers in footer
    _add_page_number_footer(doc)

    doc.save(output_path)
    print(f"Word document with layout created successfully: {output_path}")
    return output_path


# --- pdf2docx high-fidelity path ---
def create_word_with_pdf2docx(pdf_path, output_path, start=0, end=None):
    """Use pdf2docx for higher-fidelity conversion (keeps images/tables/layout)."""
    os.makedirs(os.path.dirname(output_path) or '.', exist_ok=True)
    cv = Converter(str(pdf_path))
    try:
        cv.convert(str(output_path), start=start, end=end)
    finally:
        cv.close()
    print(f"Word document (pdf2docx) created successfully: {output_path}")
    return output_path

def create_word_document(pdf_path, output_path):
    """Create Word document from PDF content"""
    
    # Extract content from PDF
    lines = extract_pdf_content(pdf_path)
    
    # Create new Word document
    doc = Document()
    
    # Set document margins (narrow margins for better layout)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Process each line and add to document
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            # Add blank line for spacing
            doc.add_paragraph()
            i += 1
            continue
        
        # Title: FORM 'A'
        if "FORM" in line and "'A'" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(14)
            run.bold = True
            i += 1
            continue
        
        # Main Heading: MEDIATION APPLICATION FORM
        if "MEDIATION APPLICATION FORM" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(16)
            run.bold = True
            i += 1
            continue
        
        # Sub-heading: [REFER RULE 3(1)]
        if "REFER RULE" in line or "[REFER" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(11)
            i += 1
            continue
        
        # Authority names
        if "Mumbai District Legal Services Authority" in line or \
           "City Civil Court, Mumbai" in line:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.bold = True
            i += 1
            continue
        
        # Section headings (all caps with colon)
        if line.isupper() and line.endswith(':'):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.bold = True
            run.underline = True
            i += 1
            continue
        
        # Numbered items (like "1 {{client_name}}")
        if line and line[0].isdigit() and len(line) > 1:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)
            i += 1
            continue
        
        # Section headings without colon but important
        if any(keyword in line for keyword in ["Name of", "Name,", "Address and contact details", 
                                                 "REGISTERED ADDRESS", "CORRESPONDENCE", 
                                                 "DETAILS OF", "Nature of disputes"]):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)
            if line.isupper():
                run.bold = True
            i += 1
            continue
        
        # Regular text with fields
        if "{{" in line or "{%" in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)
            i += 1
            continue
        
        # Labels like "Telephone No.", "Mobile No.", "Email ID", "Address"
        if any(keyword in line for keyword in ["Telephone No.", "Mobile No.", "Email ID", 
                                                 "Address", "Name", "info@kslegal.co.in",
                                                 "Applicant", "Defendant"]):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)
            i += 1
            continue
        
        # Rules and Acts
        if "COMM. COURTS" in line or "Commercial Courts Act" in line or "RULES" in line:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)
            i += 1
            continue
        
        # Default: regular paragraph
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        i += 1
    
    # Save document
    doc.save(output_path)
    print(f"Word document created successfully: {output_path}")
    return output_path

if __name__ == "__main__":
    # Input and output file paths
    pdf_file = "django_assignment.pdf"
    output_file = os.path.join("outputs", "replicated_document_layout.docx")
    
    try:
        # First try pdf2docx for maximum fidelity; fallback to our heuristic extractor
        try:
            create_word_with_pdf2docx(pdf_file, output_file)
        except Exception:
            print("pdf2docx path failed; falling back to custom layout extractor")
            create_word_document_with_layout(pdf_file, output_file)

        print(f"\n✓ Successfully created {output_file}")
        print(f"✓ The document replicates the layout and content of {pdf_file}")
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()
